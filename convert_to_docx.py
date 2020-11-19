# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn

'''
传入多篇文章dict组成的数组
存储到demo.doc中
'''
def convert_to_docx(articles):
    document = Document()

    for article in articles:
        title = article['title']
        content_arr = article['content_arr']
        document.add_heading(title, 0)
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        for content in content_arr:
            document.add_paragraph("    " + content + '\n')

        # change page
        document.add_page_break()
    
    document.save('demo.docx')